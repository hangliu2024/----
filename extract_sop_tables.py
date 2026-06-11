"""从原始docx提取SOP中的表格，生成完整的带表格HTML内容，并更新数据库"""
from docx import Document
import pymysql

# 1. 读取docx文件
doc = Document(r'd:\资产管理\案件调查sop\匿名举报案件调查SOP.docx')

print(f"文档中找到 {len(doc.tables)} 个表格")

# 2. 打印每个表格的内容
for i, table in enumerate(doc.tables):
    print(f"\n=== 表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列 ===")
    for j, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
        print(f"  行{j}: {cells}")

# 3. 将表格转换为HTML
def table_to_html(table, table_id=""):
    html = '<table class="sop-table" id="' + table_id + '">\n'
    for i, row in enumerate(table.rows):
        tag = 'th' if i == 0 else 'td'
        html += '  <tr>\n'
        for cell in row.cells:
            text = cell.text.strip().replace('\n', '<br>')
            html += '    <' + tag + '>' + text + '</' + tag + '>\n'
        html += '  </tr>\n'
    html += '</table>\n'
    return html

# 4. 生成所有表格的HTML
tables_html = []
for i, table in enumerate(doc.tables):
    html = table_to_html(table, "sop-table-" + str(i+1))
    tables_html.append(html)
    print("\n表格 " + str(i+1) + " HTML预览 (前200字符):")
    print(html[:200])

# 5. 读取当前数据库中的full_content
conn = pymysql.connect(
    host='10.5.192.253',
    port=3307,
    user='nocobase',
    password='nocobase',
    database='nocobase',
    charset='utf8mb4'
)
cursor = conn.cursor()

cursor.execute("SELECT full_content FROM investigation_sop WHERE sop_no='SOP-NJ-001'")
row = cursor.fetchone()
if row and row[0]:
    current_content = row[0]
    print("\n当前数据库中 full_content 长度: " + str(len(current_content)))
    has_table = '<table' in current_content
    print("包含 table 标签: " + str(has_table))
else:
    current_content = None
    print("\n数据库中无 full_content")

cursor.close()
conn.close()

print("\n=== 提取完成 ===")
print("共提取 " + str(len(tables_html)) + " 个表格的HTML")